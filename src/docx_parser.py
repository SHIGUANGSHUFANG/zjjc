from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Dict, Iterable, List, Optional
from io import BytesIO
import base64
import re

from docx import Document


QUIZ_SECTION_KEYWORDS = [
    "试题部分", "试题", "练习部分", "练习", "习题部分", "习题", "检测题", "检测",
    "巩固练习", "课堂练习", "课后练习", "随堂练习", "拓展训练", "综合训练",
    "例题", "真题", "选择题", "填空题", "判断题", "计算题",
]


@dataclass
class Node:
    id: str
    title: str
    level: int
    notes: List[str] = field(default_factory=list)
    images: List[Dict[str, str]] = field(default_factory=list)
    children: List["Node"] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "title": self.title,
            "level": self.level,
            "notes": self.notes,
            "images": self.images,
            "children": [c.to_dict() for c in self.children],
        }


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip("：:。. ")


def heading_level(style_name: str) -> Optional[int]:
    m = re.search(r"(?:Heading|标题)\s*(\d+)", style_name or "", re.IGNORECASE)
    return int(m.group(1)) if m else None


def is_quiz_section_title(text: str, keywords: Iterable[str] = QUIZ_SECTION_KEYWORDS) -> bool:
    t = normalize_text(text)
    if not t:
        return False
    return any(t == normalize_text(k) or t.startswith(normalize_text(k)) for k in keywords)


def paragraph_images(paragraph, document) -> List[Dict[str, str]]:
    images: List[Dict[str, str]] = []
    blips = paragraph._p.xpath(".//a:blip")
    for blip in blips:
        rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not rid or rid not in document.part.related_parts:
            continue
        image_part = document.part.related_parts[rid]
        content_type = image_part.content_type or "image/png"
        b64 = base64.b64encode(image_part.blob).decode("ascii")
        images.append({
            "data_url": f"data:{content_type};base64,{b64}",
            "content_type": content_type,
        })
    return images


def strip_blank_markup(text: str) -> str:
    # 学习阶段显示答案原文，不显示 {{}}
    return re.sub(r"\{\{\s*(.*?)\s*\}\}", r"\1", text or "")


def blank_question_text(text: str) -> str:
    return re.sub(r"\{\{\s*(.*?)\s*\}\}", "____", text or "")


def find_blanks_in_text(text: str) -> List[str]:
    return [m.group(1).strip() for m in re.finditer(r"\{\{\s*(.*?)\s*\}\}", text or "") if m.group(1).strip()]


def inspect_docx_paragraphs(file_bytes: bytes) -> List[Dict[str, str]]:
    doc = Document(BytesIO(file_bytes))
    rows = []
    for i, p in enumerate(doc.paragraphs, 1):
        text = p.text.strip()
        if not text and not paragraph_images(p, doc):
            continue
        rows.append({
            "序号": str(i),
            "样式": p.style.name if p.style else "",
            "标题级别": str(heading_level(p.style.name if p.style else "") or ""),
            "是否试题区标题": "是" if is_quiz_section_title(text) else "",
            "文字": text,
            "图片数": str(len(paragraph_images(p, doc))),
        })
    return rows


def _new_node(counter: int, title: str, level: int) -> Node:
    return Node(id=f"n{counter}", title=title.strip(), level=level)


def _append_child(stack: Dict[int, Node], node: Node, root: Optional[Node]) -> Optional[Node]:
    if node.level == 1 or root is None:
        root = node
    else:
        parent_level = node.level - 1
        while parent_level > 0 and parent_level not in stack:
            parent_level -= 1
        parent = stack.get(parent_level) or root
        parent.children.append(node)
    stack[node.level] = node
    # 清除比当前更深的历史层级，避免后续正文挂错
    for lv in list(stack.keys()):
        if lv > node.level:
            del stack[lv]
    return root


def _current_node(stack: Dict[int, Node], root: Optional[Node]) -> Optional[Node]:
    if not stack:
        return root
    return stack[max(stack.keys())]


def _parse_answer_line(text: str) -> Optional[str]:
    patterns = [
        r"^\[\s*(?:answer|答案|正确答案)\s*\]\s*[:：]?\s*([A-Da-d]|.+)$",
        r"^(?:answer|答案|正确答案)\s*[:：]\s*([A-Da-d]|.+)$",
    ]
    for pat in patterns:
        m = re.match(pat, text.strip(), re.I)
        if m:
            return m.group(1).strip()
    return None


def _parse_option_line(text: str) -> Optional[Dict[str, str]]:
    m = re.match(r"^\[?\s*([A-Da-d])\s*\]?\s*[\.、．:：]?\s*(.+)$", text.strip())
    if m:
        return {"key": m.group(1).upper(), "text": m.group(2).strip()}
    return None


def _looks_like_question_start(text: str, level: Optional[int]) -> bool:
    t = text.strip()
    if not t:
        return False
    if level is not None and level <= 4:
        return True
    if re.match(r"^\s*(\d+|[一二三四五六七八九十]+)[\.、．)]\s*", t):
        return True
    if "{{" in t and "}}" in t:
        return True
    return False


def _finalize_question(q: Optional[Dict[str, Any]], questions: List[Dict[str, Any]]) -> None:
    if not q:
        return
    prompt = "\n".join(q.get("prompt_lines", [])).strip()
    if not prompt:
        return
    blanks = find_blanks_in_text(prompt)
    qtype = "mcq" if q.get("options") else "blank"
    answer = q.get("answer") or (blanks[0] if blanks else "")
    # 如果答案是 A/B/C/D，保留字母；否则保留原文。
    questions.append({
        "id": f"q{len(questions) + 1}",
        "type": qtype,
        "prompt": prompt,
        "prompt_for_user": blank_question_text(prompt),
        "answer": answer,
        "options": q.get("options", []),
        "explanation": "\n".join(q.get("explanation", [])).strip(),
    })


def _parse_quiz_paragraph(text: str, level: Optional[int], current_q: Optional[Dict[str, Any]], questions: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not text:
        return current_q

    ans = _parse_answer_line(text)
    if ans is not None:
        if current_q is None:
            current_q = {"prompt_lines": [], "options": [], "answer": "", "explanation": []}
        current_q["answer"] = ans
        return current_q

    opt = _parse_option_line(text)
    if opt is not None:
        if current_q is None:
            current_q = {"prompt_lines": [], "options": [], "answer": "", "explanation": []}
        current_q["options"].append(opt)
        return current_q

    if _looks_like_question_start(text, level):
        _finalize_question(current_q, questions)
        return {"prompt_lines": [text], "options": [], "answer": "", "explanation": []}

    if current_q is None:
        return {"prompt_lines": [text], "options": [], "answer": "", "explanation": []}

    # 没有选项/答案标记的普通段落，作为题干续行或解析。
    if current_q.get("options") or current_q.get("answer"):
        current_q["explanation"].append(text)
    else:
        current_q["prompt_lines"].append(text)
    return current_q


def collect_blank_items(node: Dict[str, Any], trail: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    trail = (trail or []) + [strip_blank_markup(node.get("title", ""))]
    items: List[Dict[str, Any]] = []

    for idx, note in enumerate(node.get("notes", [])):
        answers = find_blanks_in_text(note)
        for answer in answers:
            items.append({
                "id": f"{node.get('id')}_note_{idx}_{len(items)}",
                "path": " / ".join([x for x in trail if x]),
                "source": note,
                "question": blank_question_text(note),
                "answer": answer,
            })

    # 标题中也允许 {{}}
    for answer in find_blanks_in_text(node.get("title", "")):
        items.append({
            "id": f"{node.get('id')}_title",
            "path": " / ".join([x for x in trail[:-1] if x]),
            "source": node.get("title", ""),
            "question": blank_question_text(node.get("title", "")),
            "answer": answer,
        })

    for child in node.get("children", []):
        items.extend(collect_blank_items(child, trail))
    return items


def parse_docx(file_bytes: bytes, h4_mode: str = "notes") -> Dict[str, Any]:
    doc = Document(BytesIO(file_bytes))
    root: Optional[Node] = None
    stack: Dict[int, Node] = {}
    counter = 0
    in_quiz_section = False
    current_q: Optional[Dict[str, Any]] = None
    quiz_questions: List[Dict[str, Any]] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        imgs = paragraph_images(p, doc)
        level = heading_level(p.style.name if p.style else "")

        # 试题区是阶段三，绝不进入导图节点/说明。
        if text and is_quiz_section_title(text):
            in_quiz_section = True
            _finalize_question(current_q, quiz_questions)
            current_q = None
            continue

        if in_quiz_section:
            if text:
                current_q = _parse_quiz_paragraph(text, level, current_q, quiz_questions)
            # 试题图片暂时作为上一题解析图片不展示，避免题目解析混乱。后续可扩展。
            continue

        if level is not None and level <= 3:
            counter += 1
            node = _new_node(counter, strip_blank_markup(text), level)
            if imgs:
                node.images.extend(imgs)
            root = _append_child(stack, node, root)
            continue

        current = _current_node(stack, root)

        if level is not None and level >= 4 and h4_mode == "nodes":
            counter += 1
            node = _new_node(counter, strip_blank_markup(text), min(level, 4))
            if imgs:
                node.images.extend(imgs)
            root = _append_child(stack, node, root)
            continue

        # Heading 4+ 默认作为说明文字；普通正文也作为说明文字。
        if text:
            if current is None:
                counter += 1
                current = _new_node(counter, "未命名主题", 1)
                root = _append_child(stack, current, root)
            current.notes.append(text)
        if imgs:
            if current is None:
                counter += 1
                current = _new_node(counter, "未命名主题", 1)
                root = _append_child(stack, current, root)
            current.images.extend(imgs)

    _finalize_question(current_q, quiz_questions)

    if root is None:
        root = Node(id="n1", title="未识别到标题", level=1, notes=["请在 Word 中使用“标题 1/2/3”样式。"])

    tree = root.to_dict()
    blanks = collect_blank_items(tree)

    return {
        "tree": tree,
        "blanks": blanks,
        "quiz_questions": quiz_questions,
        "paragraphs": inspect_docx_paragraphs(file_bytes),
    }
