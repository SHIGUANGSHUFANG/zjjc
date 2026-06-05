"""
把“标题层级 + 挖空标记 + 试题部分”的 Word 文档生成三阶段 XMind 风格 HTML 学习页。

三阶段逻辑：
1. 第一阶段：学习预览版，显示完整答案。
2. 点击“记住了”后进入第二阶段：挖空练习版。
3. 挖空提交核对后，如果正确率达到 75% 才出现“进入试题”按钮；否则必须重做。
4. 选择题提交后显示正确率、答案和解析，并允许重新做题。

Word 格式约定：
- Heading 1：中心主题
- Heading 2：主分支
- Heading 3：知识点节点
- Heading 4：知识点说明文字，可包含 [正确答案|干扰项1|干扰项2]
- 图片会挂到最近的知识点/说明节点下
- “试题部分”之后使用：
  【第X题】：
  【题目内容】：...
  【正确选项】：...
  【错误选项1】：...
  【解析】：...

运行示例：
python generate_xmind_three_stage_html.py 内能(3).docx 内能_三阶段XMind练习版.html
"""
from __future__ import annotations

import base64
import json
import re
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn


@dataclass
class Node:
    id: str
    title: str
    level: int
    notes: List[str] = field(default_factory=list)
    images: List[Dict[str, str]] = field(default_factory=list)
    children: List["Node"] = field(default_factory=list)


def heading_level(paragraph) -> Optional[int]:
    name = paragraph.style.name if paragraph.style is not None else ""
    m = re.search(r"(?:Heading|标题)\s*(\d+)", name, re.I)
    if m:
        return int(m.group(1))
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.outlineLvl is not None:
        return int(ppr.outlineLvl.val) + 1
    return None


def collect_image_map(docx_path: Path) -> Dict[str, Dict[str, str]]:
    image_map: Dict[str, Dict[str, str]] = {}
    with zipfile.ZipFile(docx_path) as zf:
        rels_path = "word/_rels/document.xml.rels"
        if rels_path not in zf.namelist():
            return image_map
        rels_xml = zf.read(rels_path).decode("utf-8", errors="ignore")
        for rid, target in re.findall(r'Id="([^"]+)"[^>]+Target="([^"]+)"', rels_xml):
            if not target.startswith("media/"):
                continue
            media_path = "word/" + target
            if media_path not in zf.namelist():
                continue
            raw = zf.read(media_path)
            suffix = Path(target).suffix.lower().lstrip(".") or "png"
            mime = "jpeg" if suffix in {"jpg", "jpeg"} else suffix
            image_map[rid] = {
                "data_url": f"data:image/{mime};base64," + base64.b64encode(raw).decode("ascii"),
                "content_type": f"image/{mime}",
                "filename": Path(target).name,
            }
    return image_map


def images_in_paragraph(paragraph, image_map: Dict[str, Dict[str, str]]) -> List[Dict[str, str]]:
    images: List[Dict[str, str]] = []
    for blip in paragraph._p.xpath('.//*[local-name()="blip"]'):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid and rid in image_map:
            images.append(image_map[rid])
    return images


def parse_docx(docx_path: Path) -> Node:
    doc = Document(docx_path)
    image_map = collect_image_map(docx_path)
    root = Node(id="n0", title=docx_path.stem, level=1)
    stack: Dict[int, Node] = {}
    last_node: Optional[Node] = None
    counter = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "试题部分":
            break
        imgs = images_in_paragraph(para, image_map)
        lvl = heading_level(para)

        if text:
            if lvl is None:
                if last_node is None:
                    root.notes.append(text)
                    last_node = root
                else:
                    last_node.notes.append(text)
                if imgs:
                    last_node.images.extend(imgs)
                continue

            if lvl == 1:
                if not root.children and not root.notes and root.title == docx_path.stem:
                    root.title = text
                    root.id = f"n{counter}"; counter += 1
                    stack = {1: root}
                    last_node = root
                else:
                    node = Node(id=f"n{counter}", title=text, level=1)
                    counter += 1
                    root.children.append(node)
                    stack[1] = node
                    last_node = node
                if imgs:
                    last_node.images.extend(imgs)
                continue

            if lvl == 4:
                parent = stack.get(3) or stack.get(2) or root
                parent.notes.append(text)
                if imgs:
                    parent.images.extend(imgs)
                last_node = parent
                continue

            parent_level = max([k for k in stack.keys() if k < lvl], default=1)
            parent = stack.get(parent_level, root)
            node = Node(id=f"n{counter}", title=text, level=lvl)
            counter += 1
            node.images.extend(imgs)
            parent.children.append(node)
            stack[lvl] = node
            for k in list(stack.keys()):
                if k > lvl:
                    del stack[k]
            last_node = node
        elif imgs:
            (last_node or root).images.extend(imgs)

    return root


QUESTION_MARK_RE = re.compile(r"^【第.+?题】[:：]?$")
FIELD_RE = re.compile(r"^【(.+?)】[:：](.*)$")


def parse_questions(docx_path: Path) -> List[Dict[str, object]]:
    """解析“试题部分”之后的选择题。"""
    doc = Document(docx_path)
    questions: List[Dict[str, object]] = []
    in_section = False
    cur: Optional[Dict[str, object]] = None
    last_field: Optional[str] = None

    def flush() -> None:
        nonlocal cur
        if cur and cur.get("stem") and cur.get("answer"):
            wrong = [x for x in cur.get("wrong", []) if str(x).strip()]
            cur["wrong"] = wrong
            cur["options"] = [cur["answer"]] + wrong
            questions.append(cur)
        cur = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text == "试题部分":
            in_section = True
            continue
        if not in_section:
            continue

        if QUESTION_MARK_RE.match(text):
            flush()
            title = text.strip("【】 ：:")
            cur = {"title": title, "stem": "", "answer": "", "wrong": [], "analysis": ""}
            last_field = None
            continue
        if cur is None:
            continue

        m = FIELD_RE.match(text)
        if m:
            key = m.group(1).replace(" ", "")
            value = m.group(2).strip()
            if key == "题目内容":
                cur["stem"] = value
                last_field = "stem"
            elif key == "正确选项":
                cur["answer"] = value
                last_field = "answer"
            elif key.startswith("错误选项"):
                cur.setdefault("wrong", []).append(value)
                last_field = "wrong"
            elif key == "解析":
                cur["analysis"] = value
                last_field = "analysis"
            else:
                last_field = None
        else:
            # 兼容 Word 自动换行或分段导致的续写。
            if last_field == "stem":
                cur["stem"] = (str(cur.get("stem", "")) + text).strip()
            elif last_field == "answer":
                cur["answer"] = (str(cur.get("answer", "")) + text).strip()
            elif last_field == "analysis":
                cur["analysis"] = (str(cur.get("analysis", "")) + text).strip()
            elif last_field == "wrong" and cur.get("wrong"):
                cur["wrong"][-1] = (str(cur["wrong"][-1]) + text).strip()

    flush()
    for i, q in enumerate(questions, 1):
        q["id"] = f"q{i}"
        q["index"] = i
    return questions


HTML_TEMPLATE = r'''<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>__TITLE__</title>
<style>
  :root {
    --bg: #f7f7fb;
    --ink: #1f2937;
    --muted: #6b7280;
    --line: #c8d0dc;
    --card: rgba(255,255,255,.96);
    --ok: #16a34a;
    --bad: #dc2626;
    --blank-bg: #eaf2ff;
    --blank-border: #9ebbf4;
  }
  html, body { margin:0; padding:0; background:var(--bg); font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","PingFang SC","Microsoft YaHei",sans-serif; }
  .toolbar { position:sticky; top:0; z-index:20; display:flex; gap:8px; align-items:center; padding:10px 14px; background:rgba(247,247,251,.92); backdrop-filter:blur(8px); border-bottom:1px solid #e5e7eb; }
  .toolbar button { border:1px solid #d1d5db; background:white; border-radius:10px; padding:7px 12px; cursor:pointer; color:#374151; box-shadow:0 1px 2px rgba(0,0,0,.05); }
  .toolbar button:hover { background:#f9fafb; }
  .toolbar .primary { background:#111827; color:white; border-color:#111827; }
  .toolbar .learn { background:linear-gradient(135deg,#10b981,#059669); color:white; border-color:#059669; font-weight:800; }
  .hint { color:#6b7280; font-size:13px; margin-left:8px; }
  #phaseBadge { font-size:13px; font-weight:800; color:#111827; background:#fff; border:1px solid #e5e7eb; border-radius:999px; padding:5px 10px; }
  #viewport { width:100%; height:calc(100vh - 52px); min-height:760px; overflow:hidden; position:relative; background:
    radial-gradient(circle at 24px 24px, rgba(99,102,241,.08) 2px, transparent 2.5px) 0 0 / 34px 34px,
    linear-gradient(180deg, #fbfbff, #f2f4f8); }
  #canvas { position:absolute; left:0; top:0; transform-origin:0 0; }
  svg { position:absolute; left:0; top:0; overflow:visible; pointer-events:none; }
  .node { position:absolute; box-sizing:border-box; background:var(--card); border:1px solid rgba(31,41,55,.10); border-radius:18px; box-shadow:0 10px 25px rgba(31,41,55,.10); overflow:hidden; }
  .node.root { background:linear-gradient(135deg,#202236,#3b3f63); color:white; text-align:center; border:none; box-shadow:0 18px 40px rgba(31,41,55,.25); }
  .node.level2 { border-top:6px solid var(--accent); }
  .node.level3 { border-left:8px solid var(--accent); }
  .node-inner { padding:14px 16px; }
  .root .node-inner { padding:20px 24px; }
  .title { font-weight:800; letter-spacing:.2px; color:#111827; line-height:1.28; }
  .root .title { color:white; font-size:24px; }
  .level2 .title { font-size:20px; color:#111827; }
  .level3 .title { font-size:17px; color:#111827; }
  .notes { margin-top:10px; display:flex; flex-direction:column; gap:8px; }
  .note { font-size:14px; line-height:1.65; color:#374151; background:#f8fafc; border:1px solid #edf1f7; border-radius:12px; padding:9px 11px; }
  .level3 .note { font-size:15px; line-height:1.72; }
  .imgs { margin-top:10px; display:flex; flex-wrap:wrap; gap:10px; align-items:flex-start; }
  .imgs img { max-width:190px; max-height:150px; object-fit:contain; background:white; border:1px solid #e5e7eb; border-radius:12px; padding:6px; box-shadow:0 4px 12px rgba(31,41,55,.10); }
  path.link { fill:none; stroke:var(--line); stroke-width:2.4; stroke-linecap:round; }
  path.link.level2 { stroke-width:3.4; stroke:var(--accent); opacity:.75; }
  .blank {
    display:inline-block; vertical-align:-0.12em; min-width:3.6em; height:1.34em; margin:0 .08em; border-radius:7px;
    background:var(--blank-bg); border:1px solid var(--blank-border); box-shadow:inset 0 1px 0 rgba(255,255,255,.95); cursor:pointer;
  }
  .blank.filled { min-width:auto; height:auto; padding:0 .38em; line-height:1.34; color:#1d4ed8; font-weight:700; background:#eff6ff; }
  .blank.correct { color:var(--ok); background:#ecfdf5; border-color:#86efac; }
  .blank.wrong { color:var(--bad); background:#fef2f2; border-color:#fecaca; }
  .answer-inline { color:var(--ok); font-weight:700; margin-left:.08em; }
  .preview-answer { color:#111827; font-weight:700; }
  #optionMenu { position:fixed; z-index:50; display:none; min-width:110px; padding:6px; border:1px solid #e5e7eb; border-radius:12px; background:white; box-shadow:0 16px 40px rgba(15,23,42,.18); }
  #optionMenu button { display:block; width:100%; text-align:left; border:0; background:white; color:#111827; border-radius:8px; padding:8px 10px; cursor:pointer; font-size:14px; }
  #optionMenu button:hover { background:#f3f4f6; }
  .score { color:#374151; font-size:13px; margin-left:8px; }

  #questionPage { display:none; min-height:calc(100vh - 52px); padding:30px 18px 80px; background:linear-gradient(180deg,#fbfbff,#f2f4f8); }
  .question-wrap { max-width:980px; margin:0 auto; }
  .question-head { background:white; border:1px solid #e5e7eb; border-radius:22px; padding:20px 24px; box-shadow:0 10px 25px rgba(31,41,55,.08); margin-bottom:18px; }
  .question-head h1 { margin:0; font-size:26px; color:#111827; }
  .question-head p { margin:8px 0 0; color:#6b7280; }
  .q-card { background:white; border:1px solid #e5e7eb; border-radius:18px; padding:18px 20px; margin:16px 0; box-shadow:0 8px 20px rgba(31,41,55,.07); }
  .q-title { font-weight:900; color:#111827; margin-bottom:8px; }
  .q-stem { font-size:16px; line-height:1.65; color:#1f2937; margin-bottom:12px; }
  .q-option { display:flex; align-items:flex-start; gap:8px; padding:10px 12px; border:1px solid #edf1f7; border-radius:12px; background:#f8fafc; margin:8px 0; cursor:pointer; line-height:1.55; }
  .q-option:hover { background:#f3f4f6; }
  .q-option input { margin-top:5px; }
  .q-option.correct-choice { border-color:#86efac; background:#ecfdf5; color:#166534; font-weight:700; }
  .q-option.wrong-choice { border-color:#fecaca; background:#fef2f2; color:#991b1b; font-weight:700; }
  .q-analysis { display:none; margin-top:12px; padding:12px 14px; border-radius:12px; background:#fff7ed; border:1px solid #fed7aa; color:#7c2d12; line-height:1.7; }
  .q-summary { display:none; margin:18px 0; padding:15px 18px; border-radius:16px; background:#111827; color:white; font-weight:900; }
  .question-actions { position:sticky; bottom:14px; display:flex; justify-content:flex-end; gap:10px; max-width:980px; margin:20px auto 0; }
  .question-actions button { border:0; border-radius:14px; padding:11px 18px; background:#111827; color:white; font-weight:800; cursor:pointer; box-shadow:0 10px 24px rgba(31,41,55,.18); }
  .question-actions button.secondary { background:white; color:#111827; border:1px solid #d1d5db; }

</style>
</head>
<body>
<div class="toolbar">
  <span id="phaseBadge">学习预览</span>
  <button onclick="zoomBy(1.15)">放大</button>
  <button onclick="zoomBy(0.87)">缩小</button>
  <button onclick="fitView()">适应窗口</button>
  <button onclick="resetView()">重置</button>
  <button id="rememberBtn" class="learn" onclick="startQuiz()">记住了</button>
  <button id="submitBtn" class="primary" onclick="submitAnswers()" style="display:none">提交挖空</button>
  <button id="nextQuestionBtn" class="learn" onclick="goQuestions()" style="display:none">进入试题</button>
  <button id="redoBtn" onclick="resetAnswers()" style="display:none">重做</button>
  <span id="score" class="score"></span>
  <span id="hint" class="hint">先看完整导图；学习结束后点击“记住了”进入挖空练习，进入后页面不提供返回预览。</span>
</div>
<div id="viewport"><div id="canvas"><svg id="links"></svg><div id="nodes"></div></div></div>
<div id="questionPage"><div class="question-wrap"><div class="question-head"><h1>试题练习</h1><p>完成全部选择题后提交，系统会显示正确率、正误和解析；核对后可以重新做题。</p></div><div id="qSummary" class="q-summary"></div><div id="questionList"></div></div><div class="question-actions"><button id="redoQuestionsBtn" class="secondary" onclick="redoQuestions()" style="display:none">重新做题</button><button id="submitQuestionsBtn" onclick="submitQuestions()">提交试题</button></div></div>
<div id="optionMenu"></div>
<script>
const tree = __TREE_JSON__;
const questions = __QUESTIONS_JSON__;
const colors = ['#EF6C6C','#F59E0B','#11A683','#5B8DEF','#8B5CF6','#14B8A6','#EC4899','#64748B'];
const levelWidths = {1: 150, 2: 130, 3: 430};
const gapX = 72;
const gapY = 24;
let mode = 'preview';
let scale = 1, tx = 20, ty = 20;
let positioned = [];
let clozeCounter = 1;

function stripCloze(s) { return (s || '').replace(/\[([^\[\]|]+(?:\|[^\[\]|]+)+)\]/g, '□□'); }
function measureNode(node) {
  const w = levelWidths[Math.min(node.level, 3)] || 360;
  const noteCount = (node.notes || []).length;
  const imgCount = (node.images || []).length;
  const titleLen = [...stripCloze(node.title || '')].length;
  let h = node.level === 1 ? 88 : 58;
  h += Math.ceil(titleLen / (node.level === 3 ? 22 : 8)) * 8;
  if (noteCount) {
    for (const note of node.notes) {
      const len = [...stripCloze(note)].length;
      h += 30 + Math.ceil(len / 30) * 20;
    }
  }
  if (imgCount) h += 175;
  return {w, h: Math.max(h, node.level === 1 ? 86 : 60)};
}
function annotate(node, depth=0, branch=0) {
  node.depth = depth;
  node.branch = depth <= 1 ? branch : branch;
  const m = measureNode(node); node.w = m.w; node.h = m.h;
  (node.children || []).forEach((c, i) => annotate(c, depth + 1, depth === 0 ? i : node.branch));
}
function layout(node, x=0, y=0) {
  const children = node.children || [];
  node.x = x;
  if (!children.length) { node.subtreeH = node.h; node.y = y; return node.subtreeH; }
  let curY = y;
  for (const c of children) { layout(c, x + node.w + gapX, curY); curY += c.subtreeH + gapY; }
  const childrenH = curY - y - gapY;
  node.subtreeH = Math.max(node.h, childrenH);
  const first = children[0], last = children[children.length - 1];
  node.y = (first.y + last.y + last.h - node.h) / 2;
  if (node.subtreeH > childrenH) {
    const diff = (node.subtreeH - childrenH) / 2;
    for (const c of children) shift(c, diff);
  }
  return node.subtreeH;
}
function shift(node, dy) { node.y += dy; (node.children || []).forEach(c => shift(c, dy)); }
function collect(node, arr=[]) { arr.push(node); (node.children || []).forEach(c => collect(c, arr)); return arr; }
function escapeHTML(s) { return (s || '').replace(/[&<>"']/g, m => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#039;'}[m])); }
function shuffle(arr) { const a = arr.slice(); for (let i=a.length-1;i>0;i--){ const j=Math.floor(Math.random()*(i+1)); [a[i],a[j]]=[a[j],a[i]]; } return a; }
function renderClozeText(s) {
  let out = '', last = 0;
  const re = /\[([^\[\]|]+(?:\|[^\[\]|]+)+)\]/g;
  let m;
  while ((m = re.exec(s || '')) !== null) {
    out += escapeHTML(s.slice(last, m.index));
    const options = m[1].split('|').map(x => x.trim()).filter(Boolean);
    const answer = options[0] || '';
    if (mode === 'preview') {
      out += `<span class="preview-answer">${escapeHTML(answer)}</span>`;
    } else {
      const id = 'c' + (clozeCounter++);
      out += `<span class="blank" id="${id}" data-answer="${escapeHTML(answer)}" data-options='${escapeHTML(JSON.stringify(options))}' tabindex="0"></span>`;
    }
    last = re.lastIndex;
  }
  out += escapeHTML((s || '').slice(last));
  return out;
}
function updateToolbar(){
  const preview = mode === 'preview';
  const quiz = mode === 'quiz';
  const questionMode = mode === 'questions';
  document.getElementById('phaseBadge').textContent = preview ? '学习预览' : (quiz ? '挖空练习' : '试题练习');
  document.getElementById('rememberBtn').style.display = preview ? '' : 'none';
  document.getElementById('submitBtn').style.display = quiz ? '' : 'none';
  document.getElementById('redoBtn').style.display = quiz ? '' : 'none';
  if (!quiz) document.getElementById('nextQuestionBtn').style.display = 'none';
  document.getElementById('score').textContent = '';
  document.getElementById('hint').textContent = preview ? '先看完整导图；学习结束后点击“记住了”进入挖空练习。' : (quiz ? '点击浅色空格块选择答案；提交核对后才会出现“进入试题”。' : '完成全部选择题后提交，系统会显示答案和解析。');
}
function render(keepView=false) {
  document.getElementById('viewport').style.display = '';
  document.getElementById('questionPage').style.display = 'none';
  clozeCounter = 1;
  annotate(tree); layout(tree, 40, 40); positioned = collect(tree);
  const maxX = Math.max(...positioned.map(n => n.x + n.w)) + 80;
  const maxY = Math.max(...positioned.map(n => n.y + n.h)) + 80;
  const canvas = document.getElementById('canvas');
  canvas.style.width = maxX + 'px'; canvas.style.height = maxY + 'px';
  const svg = document.getElementById('links');
  svg.setAttribute('width', maxX); svg.setAttribute('height', maxY); svg.innerHTML = '';
  const nodes = document.getElementById('nodes'); nodes.innerHTML = '';

  for (const n of positioned) {
    const accent = colors[n.branch % colors.length];
    for (const c of (n.children || [])) {
      const x1 = n.x + n.w, y1 = n.y + n.h/2, x2 = c.x, y2 = c.y + c.h/2;
      const mx = (x1 + x2) / 2;
      const p = document.createElementNS('http://www.w3.org/2000/svg', 'path');
      p.setAttribute('d', `M ${x1} ${y1} C ${mx} ${y1}, ${mx} ${y2}, ${x2} ${y2}`);
      p.setAttribute('class', 'link ' + (c.level === 2 ? 'level2' : ''));
      p.style.setProperty('--accent', colors[c.branch % colors.length]);
      svg.appendChild(p);
    }
    const el = document.createElement('div');
    el.className = 'node ' + (n.level === 1 ? 'root' : n.level === 2 ? 'level2' : 'level3');
    el.style.left = n.x + 'px'; el.style.top = n.y + 'px'; el.style.width = n.w + 'px'; el.style.minHeight = n.h + 'px';
    el.style.setProperty('--accent', accent);
    const notes = (n.notes || []).map(t => `<div class="note">${renderClozeText(t)}</div>`).join('');
    const imgs = (n.images || []).map(img => `<img src="${img.data_url}" alt="插图"/>`).join('');
    el.innerHTML = `<div class="node-inner"><div class="title">${renderClozeText(n.title)}</div>${notes ? `<div class="notes">${notes}</div>` : ''}${imgs ? `<div class="imgs">${imgs}</div>` : ''}</div>`;
    nodes.appendChild(el);
  }
  updateToolbar();
  if (mode === 'quiz') bindBlanks();
  if (!keepView) fitView(); else applyTransform();
}
function startQuiz(){
  mode = 'quiz';
  try { history.replaceState({phase:'quiz'}, '', location.href); history.pushState({phase:'quiz'}, '', location.href); } catch(e) {}
  render(false);
}
window.addEventListener('popstate', () => {
  if (mode === 'quiz') {
    try { history.pushState({phase:'quiz'}, '', location.href); } catch(e) {}
  }
});
function bindBlanks(){
  document.querySelectorAll('.blank').forEach(b => {
    b.addEventListener('click', e => { e.stopPropagation(); openOptions(b); });
    b.addEventListener('keydown', e => { if(e.key==='Enter' || e.key===' '){ e.preventDefault(); openOptions(b); } });
  });
}
function openOptions(blank){
  const menu = document.getElementById('optionMenu');
  const options = shuffle(JSON.parse(blank.dataset.options || '[]'));
  menu.innerHTML = options.map(o => `<button type="button" data-v="${escapeHTML(o)}">${escapeHTML(o)}</button>`).join('');
  menu.querySelectorAll('button').forEach(btn => btn.onclick = () => chooseOption(blank, btn.dataset.v));
  const r = blank.getBoundingClientRect();
  menu.style.left = Math.min(r.left, window.innerWidth - 140) + 'px';
  menu.style.top = (r.bottom + 6) + 'px';
  menu.style.display = 'block';
}
function chooseOption(blank, value){
  blank.dataset.value = value;
  blank.textContent = value;
  blank.classList.add('filled');
  blank.classList.remove('correct','wrong');
  const nxt = blank.nextElementSibling;
  if (nxt && nxt.classList && nxt.classList.contains('answer-inline')) nxt.remove();
  document.getElementById('optionMenu').style.display = 'none';
}
document.addEventListener('click', () => document.getElementById('optionMenu').style.display = 'none');
function submitAnswers(){
  let total=0, right=0;
  document.querySelectorAll('.blank').forEach(b => {
    total++;
    const value = b.dataset.value || '';
    const answer = b.dataset.answer || '';
    const old = b.nextElementSibling;
    if (old && old.classList && old.classList.contains('answer-inline')) old.remove();
    b.classList.add('filled');
    if (value === answer) { b.classList.add('correct'); b.classList.remove('wrong'); right++; }
    else {
      b.classList.add('wrong'); b.classList.remove('correct');
      if (!value) b.textContent = '　';
      const span = document.createElement('span');
      span.className = 'answer-inline';
      span.textContent = `（${answer}）`;
      b.insertAdjacentElement('afterend', span);
    }
  });
  const rate = total ? right / total : 0;
  const pct = Math.round(rate * 100);
  const scoreEl = document.getElementById('score');
  scoreEl.textContent = total ? `挖空得分：${right}/${total}（正确率：${pct}%）` : '';
  const nextBtn = document.getElementById('nextQuestionBtn');
  if (total && questions.length && rate >= 0.75) {
    nextBtn.style.display = '';
    document.getElementById('hint').textContent = '挖空正确率已达标，可以进入试题。';
  } else if (total && questions.length) {
    nextBtn.style.display = 'none';
    document.getElementById('hint').textContent = '挖空正确率低于 75%，请点击“重做”，达到 75% 后才能进入试题。';
    alert(`本次挖空正确率为 ${pct}%，低于 75%。请重做挖空练习，达标后才能进入试题。`);
  }
}

function resetAnswers(){
  document.querySelectorAll('.answer-inline').forEach(e => e.remove());
  document.querySelectorAll('.blank').forEach(b => { b.dataset.value=''; b.textContent=''; b.className='blank'; });
  document.getElementById('score').textContent='';
  document.getElementById('nextQuestionBtn').style.display='none';
}


function renderQuestions(){
  const list = document.getElementById('questionList');
  const summary = document.getElementById('qSummary');
  summary.style.display = 'none';
  summary.textContent = '';
  document.getElementById('redoQuestionsBtn').style.display = 'none';
  document.getElementById('submitQuestionsBtn').style.display = '';
  const letters = ['A','B','C','D','E','F'];
  list.innerHTML = questions.map((q, qi) => {
    const opts = shuffle(q.options || []).map((text, oi) => ({text, letter: letters[oi] || String(oi+1)}));
    return `<div class="q-card" data-answer="${escapeHTML(q.answer || '')}">
      <div class="q-title">${escapeHTML(q.title || ('第' + (qi+1) + '题'))}</div>
      <div class="q-stem">${escapeHTML(q.stem || '')}</div>
      <div class="q-options">
        ${opts.map(o => `<label class="q-option" data-value="${escapeHTML(o.text)}"><input type="radio" name="q${qi}" value="${escapeHTML(o.text)}"><span>${o.letter}. ${escapeHTML(o.text)}</span></label>`).join('')}
      </div>
      <div class="q-analysis"><b>解析：</b>${escapeHTML(q.analysis || '暂无解析')}</div>
    </div>`;
  }).join('');
}
function goQuestions(){
  mode = 'questions';
  document.getElementById('viewport').style.display = 'none';
  document.getElementById('questionPage').style.display = 'block';
  document.getElementById('optionMenu').style.display = 'none';
  renderQuestions();
  updateToolbar();
  window.scrollTo(0,0);
}
function submitQuestions(){
  const cards = Array.from(document.querySelectorAll('.q-card'));
  const missing = cards.filter((card, i) => !card.querySelector(`input[name="q${i}"]:checked`));
  if (missing.length) { alert(`还有 ${missing.length} 道题没有作答。`); return; }
  let right = 0;
  cards.forEach((card, i) => {
    const answer = card.dataset.answer || '';
    const checked = card.querySelector(`input[name="q${i}"]:checked`);
    const value = checked ? checked.value : '';
    card.querySelectorAll('.q-option').forEach(label => {
      label.classList.remove('correct-choice','wrong-choice');
      const v = label.dataset.value || '';
      if (v === answer) label.classList.add('correct-choice');
      if (v === value && v !== answer) label.classList.add('wrong-choice');
      const input = label.querySelector('input');
      if (input) input.disabled = true;
    });
    if (value === answer) right++;
    const analysis = card.querySelector('.q-analysis');
    analysis.style.display = 'block';
    analysis.innerHTML = `<b>正确答案：</b>${escapeHTML(answer)}<br><b>解析：</b>${escapeHTML((questions[i] && questions[i].analysis) || '暂无解析')}`;
  });
  const pct = cards.length ? Math.round(right / cards.length * 100) : 0;
  const summary = document.getElementById('qSummary');
  summary.style.display = 'block';
  summary.textContent = `试题得分：${right}/${cards.length}（正确率：${pct}%）`;
  document.getElementById('score').textContent = `试题得分：${right}/${cards.length}（正确率：${pct}%）`;
  document.getElementById('redoQuestionsBtn').style.display = '';
  document.getElementById('submitQuestionsBtn').style.display = 'none';
  window.scrollTo(0,0);
}

function redoQuestions(){
  renderQuestions();
  document.getElementById('score').textContent = '';
  document.getElementById('hint').textContent = '已重新生成试题选项，请完成全部选择题后再次提交。';
  window.scrollTo(0,0);
}

function applyTransform() { document.getElementById('canvas').style.transform = `translate(${tx}px, ${ty}px) scale(${scale})`; }
function fitView() {
  const vp = document.getElementById('viewport');
  const canvas = document.getElementById('canvas');
  const cw = parseFloat(canvas.style.width), ch = parseFloat(canvas.style.height);
  scale = Math.min((vp.clientWidth - 40) / cw, (vp.clientHeight - 40) / ch, 1.25);
  tx = Math.max(20, (vp.clientWidth - cw * scale) / 2);
  ty = Math.max(20, (vp.clientHeight - ch * scale) / 2);
  applyTransform();
}
function resetView() { scale = 1; tx = 20; ty = 20; applyTransform(); }
function zoomBy(k) { scale = Math.max(.2, Math.min(2.5, scale * k)); applyTransform(); }
let dragging=false, sx=0, sy=0, ox=0, oy=0;
const vp = document.getElementById('viewport');
vp.addEventListener('mousedown', e => { if(e.target.closest('.blank') || e.target.closest('#optionMenu') || e.target.closest('.toolbar')) return; dragging=true; sx=e.clientX; sy=e.clientY; ox=tx; oy=ty; });
window.addEventListener('mousemove', e => { if(!dragging) return; tx=ox+e.clientX-sx; ty=oy+e.clientY-sy; applyTransform(); });
window.addEventListener('mouseup', () => dragging=false);
vp.addEventListener('wheel', e => { e.preventDefault(); zoomBy(e.deltaY < 0 ? 1.08 : .92); }, {passive:false});
window.addEventListener('resize', fitView);
render();
</script>
</body>
</html>
'''


def build_html(root: Node, questions: List[Dict[str, object]], title: str) -> str:
    def to_dict(n: Node) -> Dict:
        return {
            "id": n.id,
            "title": n.title,
            "level": n.level,
            "notes": n.notes,
            "images": n.images,
            "children": [to_dict(c) for c in n.children],
        }
    tree_json = json.dumps(to_dict(root), ensure_ascii=False)
    questions_json = json.dumps(questions, ensure_ascii=False)
    return HTML_TEMPLATE.replace("__TITLE__", title).replace("__TREE_JSON__", tree_json).replace("__QUESTIONS_JSON__", questions_json)


def main() -> None:
    if len(sys.argv) < 3:
        print("用法：python generate_xmind_three_stage_html.py 输入.docx 输出.html")
        raise SystemExit(1)
    docx_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    root = parse_docx(docx_path)
    questions = parse_questions(docx_path)
    out_path.write_text(build_html(root, questions, out_path.stem), encoding="utf-8")
    print(f"已生成：{out_path}")


if __name__ == "__main__":
    main()
